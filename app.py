import os
import io
import nltk
import tempfile
import re
import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from google.cloud import vision_v1
from docx import Document  # Import library for DOCX file handling
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from fpdf import FPDF

# Import load_dotenv
from dotenv import load_dotenv

# Load environment variables from the .env file
load_dotenv()

nltk.download('stopwords')
nltk.download('punkt')

# --- Configure the Gemini API ---
# Fetch the API key securely from the environment variables
gemini_api_key = os.getenv("GEMINI_API_KEY")
if not gemini_api_key:
    st.error("Gemini API key not found. Please check your .env file.")

genai.configure(api_key=gemini_api_key)

# --- Set environment variable for Google Cloud credentials ---
script_dir = os.path.dirname(os.path.abspath(__file__))
credential_path = os.path.join(script_dir, "chat-and-pdf-file-project-41797fb97018.json")
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = credential_path

# --- Get response from Gemini ---
def get_response_from_gemini(context, user_query, response_length):
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        prompt = f"{context}\nUser query: {user_query}\nResponse length: {response_length}"
        response = model.generate_content(prompt)
        return response.text if response else "No response generated."
    except Exception as e:
        st.error(f"Error with Gemini API: {e}")
        return "Error generating response."

# --- Clean extracted text ---
def clean_text(text):
    text = re.sub(r'\s+', ' ', text)  # Replace multiple whitespaces with a single space
    return text.strip()  # Remove leading and trailing whitespace

# --- Preview the uploaded PDF in the sidebar ---
def preview_pdf(file):
    try:
        pdf_reader = PdfReader(file)
        full_text = "".join(page.extract_text() for page in pdf_reader.pages)
        cleaned_text = clean_text(full_text)
        st.sidebar.subheader("PDF Preview")
        st.sidebar.text_area("Content", value=cleaned_text, height=300)
    except Exception as e:
        st.sidebar.error(f"Error previewing PDF: {e}")

# --- Preview the uploaded DOCX in the sidebar ---
def preview_docx(file):
    try:
        doc = Document(file)
        full_text = "\n".join(para.text for para in doc.paragraphs)
        cleaned_text = clean_text(full_text)
        st.sidebar.subheader("DOCX Preview")
        st.sidebar.text_area("Content", value=cleaned_text, height=300)
    except Exception as e:
        st.sidebar.error(f"Error previewing DOCX: {e}")

# --- Export chat history as text ---
def export_chat_history():
    chat_history = "\n".join(
        [f"{msg['role'].capitalize()}: {msg['content']}" for msg in st.session_state.messages]
    )
    buf = io.BytesIO()
    buf.write(chat_history.encode())
    buf.seek(0)
    st.sidebar.download_button(
        label="Download Chat History as Text",
        data=buf,
        file_name="chat_history.txt",
        mime="text/plain"
    )

# --- Extract text from PDF using Vision API ---
def extract_text_from_pdf(file):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(file.read())
            temp_pdf_path = temp_pdf.name

        # Try Vision API
        client = vision_v1.ImageAnnotatorClient()
        with open(temp_pdf_path, "rb") as pdf_file:
            content = pdf_file.read()
            image = vision_v1.Image(content=content)
            response = client.document_text_detection(image=image)

        os.remove(temp_pdf_path)

        if response and response.full_text_annotation and response.full_text_annotation.text:
            text = response.full_text_annotation.text
            return [(1, clean_text(text))]        
        
        pdf_reader = PdfReader(file)
        full_text = "".join(page.extract_text() for page in pdf_reader.pages if page.extract_text())
        if full_text.strip():
            return [(1, clean_text(full_text))]
        return [(1, "No text detected in the PDF. It might be blank or unsupported.")]
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return [(1, f"An error occurred: {e}")]

# --- Extract text from DOCX ---
def extract_text_from_docx(file):
    try:
        doc = Document(file)
        sections = []
        section = ""
        for idx, paragraph in enumerate(doc.paragraphs, 1):
            section += paragraph.text + "\n"
            if idx % 5 == 0 or idx == len(doc.paragraphs):
                sections.append((f"Section {len(sections)+1}", section.strip()))
                section = ""
        return sections
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return []

# --- Check if query is related to the document ---
def is_query_related_to_document(query, doc_context):
    query_tokens = set(word_tokenize(query.lower()))
    doc_tokens = set(word_tokenize(" ".join([content for _, content in doc_context]).lower()))
    common_tokens = query_tokens.intersection(doc_tokens)
    return len(common_tokens) > 0

def is_valid_query(query):
    stop_words = set(stopwords.words('english'))
    words = query.lower().split()  # Tokenize the input
    meaningful_words = [word for word in words if word not in stop_words]  # Filter out stopwords
    return len(meaningful_words) > 0  # Query is valid if it contains meaningful words

# --- Main function ---
def main():
    st.title('Agus Robot')
    st.sidebar.header("Upload and Preview")

    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'doc_context' not in st.session_state:
        st.session_state.doc_context = []
    if 'response_length' not in st.session_state:
        st.session_state.response_length = "Medium"
    if 'user_queries' not in st.session_state:
        st.session_state.user_queries = []

    uploaded_file = st.sidebar.file_uploader("Upload a PDF or DOCX file", type=["pdf", "docx"])

    if uploaded_file:
        file_type = uploaded_file.name.split('.')[-1].lower()
        if file_type == "pdf":
            st.session_state.doc_context = extract_text_from_pdf(uploaded_file)
            st.sidebar.text_area("Extracted PDF Context", value="\n".join([text for _, text in st.session_state.doc_context]), height=300)
        elif file_type == "docx":
            st.session_state.doc_context = extract_text_from_docx(uploaded_file)
            st.sidebar.text_area("Extracted DOCX Context", value="\n".join([text for _, text in st.session_state.doc_context]), height=300)
        else:
            st.sidebar.error("Unsupported file format.")

    response_length = st.sidebar.radio(
        "Select response length", ["Short", "Medium", "Long"]
    )

    if response_length != st.session_state.response_length:
        # Step 1: Clear all messages in session state
        st.session_state.messages = []

        # Step 2: Update the response length in session state
        st.session_state.response_length = response_length

        # Step 3: Get previous user queries
        previous_queries = st.session_state.user_queries.copy()

        # Step 4: Request new responses for each previous query
        context = "\n".join([content for _, content in st.session_state.doc_context])

        for query in previous_queries:
            # Check if the query already exists in messages to prevent duplication
            if not any(msg['content'] == query for msg in st.session_state.messages):
                # Add the user query to messages if it doesn't exist
                st.session_state.messages.append({'role': 'user', 'content': query})

                # Get new response based on updated response length
                response = get_response_from_gemini(context, query, response_length)

                # Add the new response to messages
                st.session_state.messages.append({'role': 'assistant', 'content': response})

        # Step 5: Reset the user queries list
        st.session_state.user_queries = previous_queries

    st.header("Chat")
    st.markdown("Ask questions or provide prompts based on the uploaded document.")

    # Display chat messages in the chat area
    for message in st.session_state.messages:
        if message['role'] == 'user':
            st.chat_message('user').markdown(message['content'])
            st.session_state.user_queries.append(message['content'])
        else:
            st.chat_message('assistant').markdown(message['content'])

    prompt = st.chat_input("Ask a question or give a prompt:")

    if prompt:        
        # Validate query
        if not is_valid_query(prompt):
            st.warning("Please provide a more meaningful query.")
            return
       
        # Check if document context is available and has content
        if st.session_state.doc_context:  
            # If document context exists, check if the query is related
            if not is_query_related_to_document(prompt, st.session_state.doc_context):
                st.warning("Your query does not seem related to the uploaded document. Please try again.")
                return
                
        st.chat_message('user').markdown(prompt)        
        st.session_state.messages.append({'role': 'user', 'content': prompt})        
        st.session_state.user_queries.append(prompt)

        context = "\n".join([content for _, content in st.session_state.doc_context])
        response = get_response_from_gemini(context, prompt, st.session_state.response_length)

        st.chat_message('assistant').markdown(response)
        st.session_state.messages.append({'role': 'assistant', 'content': response})        

    export_chat_history()

if __name__ == "__main__":
    main()

